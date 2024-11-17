from requete_mistral import *
from docx import Document
import tkinter as tk
from tkinter import filedialog
from hybrid_search import *
from tkinter import ttk

out_put_path="C:/Users/Utilisateur/Downloads/lettre_jurydique.docx"


def save_letter_to_word(letter_text, output_path):
    document = Document()
    document.add_paragraph(letter_text)
    document.save(output_path)
    print('fichier bien enregistré')


# Fonction qui sera exécutée lorsque le bouton est cliqué
def envoyer():
    texte = entree_texte.get("1.0", tk.END) 
    liste_texte_loi=hybrid_search(texte)
    texte_final=make_final_prompt(liste_texte_loi, texte)
    reponse=appel_mistral(texte_final)
    prompt = "Je vais te donner un courrier d'avocat, j'ai besoin que tu m'extraies toutes les citations juridiques de cette lettre. Voici la lettre :\n" + reponse + "\n je veux que tu me le fasses absolument sous le format suivant, c'est très important: [FONDEMENT]référence n°1[/FONDEMENT][FONDEMENT]référence n°2[/FONDEMENT][FONDEMENT]référence n°3[/FONDEMENT]etc..."
    fondements = appel_mistral(prompt)
    fondements_liste = []
    for f in fondements.split("[FONDEMENT]"):
        if "[/FONDEMENT]" in f:
            fondements_liste.append(f.split("[/FONDEMENT]")[0] )
    afficher_texte(reponse)
    afficher_fondement(fondements_liste)

def re_run():
     
    selection = [element for element, var in valeurs_cases.items() if var.get()]
    # Vous pouvez ici utiliser 'selection' selon vos besoins (supprimer, afficher, etc.)
    print(selection)
    texte_actu=zone_affichage.get("1.0", tk.END).strip()
    print(texte_actu)

    reponse=enlever_fondement_nul(texte_actu, selection)
    print('reponse_new =: ', reponse)
    prompt = "Je vais te donner un courrier d'avocat, j'ai besoin que tu m'extraies toutes les citations juridiques (uniquement les articles et jurisprudences, pas les textes en tant que tel stp, JE VEUX UNIQUEMENT LA RÉFÉRENCE) de cette lettre. Voici la lettre :\n" + texte_actu + "\n je veux que tu me le fasses absolument sous le format suivant, c'est très important: [FONDEMENT]référence n°1[/FONDEMENT][FONDEMENT]référence n°2[/FONDEMENT][FONDEMENT]référence n°3[/FONDEMENT]etc..."
    fondements = appel_mistral(prompt)
    fondements_liste = []
    for f in fondements.split("[FONDEMENT]"):
        if "[/FONDEMENT]" in f:
            fondements_liste.append(f.split("[/FONDEMENT]")[0] )
    for widget in cadre_cases_parent.winfo_children():
        widget.destroy()
    afficher_texte(reponse)
    afficher_fondement(fondements_liste)

# Fonction pour afficher un texte donné
def afficher_texte(texte_a_afficher):
    # Supprime le contenu précédent dans la zone d'affichage
    zone_affichage.delete("1.0", tk.END)
    # Affiche le texte fourni
    zone_affichage.insert(tk.END, texte_a_afficher)


def afficher_fondement(liste_elements):
    global cadre_cases_parent, valeurs_cases
    """
    Affiche les cases à cocher dans le cadre fourni.
    
    Args:
        liste_elements (list[str]): Liste de fondements juridiques à afficher.
        cadre_parent (tk.Frame): Cadre où les cases à cocher seront placées.
    """
    # Effacer les widgets existants dans le cadre_parent pour éviter le chevauchement
    for widget in cadre_cases_parent.winfo_children():
        widget.destroy()

    # Dictionnaire pour stocker les valeurs des cases à cocher
    valeurs_cases = {}

    # Création des cases à cocher
    for idx, element in enumerate(liste_elements):
        # Variable associée à chaque case
        var = tk.BooleanVar()
        valeurs_cases[element] = var
        # Ajout d'une case à cocher
        case = ttk.Checkbutton(cadre_cases_parent, text=element, variable=var)
        case.grid(row=idx, column=0, sticky='w', padx=(0, 5), pady=2)

    # Fonction pour afficher les éléments sélectionnés
    def afficher_selection():
        selection = [element for element, var in valeurs_cases.items() if var.get()]
        print("Éléments sélectionnés :", selection)

    # Bouton pour valider les choix
    bouton_valider = ttk.Button(cadre_cases_parent, text="enlever", command=re_run)
    bouton_valider.grid(row=len(liste_elements), column=0, sticky='w', padx=5, pady=10)
    print('afficher_fondements a run ')


def enregistrer():
# Récupère le texte dans la zone d'affichage
    texte_a_sauvegarder = zone_affichage.get("1.0", tk.END).strip()
    if not texte_a_sauvegarder:
        return  # Ne rien faire si la zone d'affichage est vide

    # Ouvre une boîte de dialogue pour choisir l'emplacement du fichier
    fichier_path = filedialog.asksaveasfilename(
        defaultextension=".txt",
        filetypes=[("Fichiers texte", "*.txt"), ("Tous les fichiers", "*.*")]
    )
    if fichier_path:  # Si un fichier est sélectionné
        document = Document()
        document.add_paragraph(texte_a_sauvegarder)
        document.save(fichier_path)


if __name__ == '__main__':

    #pb= Mon propriétaire indique dans mon contrat de location que le salon fait partie de la partie privée de ma colocataire et qu'il n'y a pas de partie commune

    fenetre = tk.Tk()
    fenetre.title("Interface Graphique Fusionnée")

    # Définition du style avec ttk
    style = ttk.Style(fenetre)
    style.theme_use("clam")  # Vous pouvez changer le thème si vous le souhaitez

    # Cadre principal
    main_frame = ttk.Frame(fenetre, padding=(20, 10))
    main_frame.pack(fill=tk.BOTH, expand=True)

    # Création d'une zone de texte pour saisir le problème
    entree_label = ttk.Label(main_frame, text="Quel est votre problème ?", font=("Helvetica", 12))
    entree_label.pack(anchor='w', pady=(0, 5))

    entree_texte = tk.Text(main_frame, height=5, width=80, font=("Helvetica", 11))
    entree_texte.pack(pady=(0, 10), fill=tk.BOTH, expand=True)

    # Cadre pour organiser les boutons horizontalement
    cadre_boutons = ttk.Frame(main_frame)
    cadre_boutons.pack(pady=10)

    # Bouton "Envoyer"
    bouton_envoyer = ttk.Button(cadre_boutons, text="Envoyer", command=envoyer)
    bouton_envoyer.pack(side=tk.LEFT, padx=5)

    # Bouton "Download"
    bouton_download = ttk.Button(cadre_boutons, text="Download", command=enregistrer)
    bouton_download.pack(side=tk.LEFT, padx=5)

    # Zone pour afficher le texte fourni
    label_affichage = ttk.Label(main_frame, text="Lettre :", font=("Helvetica", 12))
    label_affichage.pack(anchor='w', pady=(20, 5))

    zone_affichage = tk.Text(main_frame, height=10, width=80, font=("Helvetica", 11))
    zone_affichage.config(state='normal', bg='#f0f0f0')
    zone_affichage.pack(pady=(0, 10), fill=tk.BOTH, expand=True)

    # Cadre pour les fondements juridiques
    fondements_frame = ttk.Frame(main_frame)
    fondements_frame.pack(fill=tk.BOTH, expand=True, pady=(20, 5))

    # Label pour "Fondements Juridiques"
    label_fondements = ttk.Label(fondements_frame, text="Fondements Juridiques :", font=("Helvetica", 12))
    label_fondements.grid(row=0, column=0, sticky='nw', padx=(0, 10))

    # Cadre pour les cases à cocher
    cadre_cases_parent = ttk.Frame(fondements_frame, padding=(10, 10))
    cadre_cases_parent.grid(row=0, column=1, sticky='nw')

    # Ajustement de la taille minimale de la fenêtre
    fenetre.minsize(800, 600)

    # Centrer la fenêtre sur l'écran
    fenetre.update_idletasks()
    width = fenetre.winfo_width()
    height = fenetre.winfo_height()
    x = (fenetre.winfo_screenwidth() // 2) - (width // 2)
    y = (fenetre.winfo_screenheight() // 2) - (height // 2)
    fenetre.geometry(f'{width}x{height}+{x}+{y}')

    # Lancement de la boucle principale de l'interface graphique
    fenetre.mainloop()
    